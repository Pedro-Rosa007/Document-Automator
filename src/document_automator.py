import pandas as pd
import win32com.client as win32
import os
import sys
import traceback
import json
import re
import time
from datetime import datetime
from collections import defaultdict
import pythoncom
import unicodedata
import getpass
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import io

# ===============================
# CONFIGURAÇÕES INICIAIS
# ===============================

# Configuração centralizada
CONFIG = {
    "diretorios": {},
    "placeholders": {},
    "config_geral": {
        "executar_em_segundo_plano": True,
        "max_workers": 4,
        "tentativas": 3,
        "intervalo_tentativas": 2,
        "padrao_nome_arquivo": "Documento_[CONTADOR].docx"
    },
    "organizacao": {},
    "pre_pos_processamento": {}
}

# Obter caminho seguro para o arquivo de configuração
def get_config_path():
    # Usar a pasta de documentos do usuário
    user_dir = os.path.join(os.path.expanduser('~'), 'Documents')
    config_dir = os.path.join(user_dir, 'DocumentAutomator')
    
    # Criar diretório se não existir
    if not os.path.exists(config_dir):
        try:
            os.makedirs(config_dir)
        except Exception as e:
            print(f"⚠ Não foi possível criar diretório de configuração: {str(e)}")
            # Usar diretório atual como fallback
            return 'doc_automator_config.json'
    
    return os.path.join(config_dir, 'doc_automator_config.json')

CONFIG_FILE = get_config_path()

# Função para limpar caminhos
def limpar_caminho(caminho):
    """Remove aspas e normaliza caminhos"""
    try:
        if not caminho:
            return caminho
            
        # Remover aspas exteriores
        caminho = caminho.strip('"').strip("'")
        
        # Substituir barras invertidas duplas por simples
        caminho = caminho.replace('\\\\', '\\')
        
        # Normalizar caminho
        return os.path.normpath(caminho)
    except Exception as e:
        print(f"Erro ao normalizar caminho: {str(e)}")
        return caminho

# ===============================
# FUNÇÕES AUXILIARES
# ===============================

def salvar_configuracao():
    try:
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(CONFIG, f, ensure_ascii=False, indent=4)
        print(f"✓ Configuração salva com sucesso em: {CONFIG_FILE}")
    except Exception as e:
        print(f"⚠ Erro ao salvar configuração: {str(e)}")
        print("Verifique as permissões do diretório.")

def carregar_configuracao():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"⚠ Erro ao carregar configuração: {str(e)}")
            print("Iniciando com configuração padrão.")
    return {}

def perguntar_sim_nao(prompt, padrao=None):
    while True:
        resposta = input(prompt + " [s/n] ").strip().lower()
        if resposta in ['s', 'n']:
            return resposta == 's'
        if padrao is not None:
            return padrao
        print("⚠ Responda com 's' ou 'n'")

def perguntar_path(prompt, deve_existir=True, padrao=None):
    while True:
        try:
            caminho = input(prompt).strip()
            if padrao and not caminho:
                caminho = padrao
                
            # Aplicar limpeza ao caminho
            caminho = limpar_caminho(caminho)
            
            if not caminho:
                print("⚠ Este campo é obrigatório!")
                continue
                
            # Verificar se caminho existe
            if deve_existir:
                if not os.path.exists(caminho):
                    print(f"⚠ Caminho não existe: {caminho}")
                    print("Verifique o caminho e tente novamente.")
                    continue
                elif not os.path.isdir(caminho) and not os.path.isfile(caminho):
                    print(f"⚠ Caminho não é um diretório ou arquivo válido: {caminho}")
                    continue
                    
            return caminho
        except Exception as e:
            print(f"Erro: {str(e)} - Tente novamente")

def limpar_nome_arquivo(nome):
    """Remove caracteres inválidos de nomes de arquivos/pastas"""
    nome = unicodedata.normalize('NFKD', nome).encode('ascii', 'ignore').decode('ascii')
    nome = re.sub(r'[\\/*?:"<>|]', "_", nome)
    nome = re.sub(r'\s+', ' ', nome).strip()
    return nome[:100]  # Limitar tamanho

# ===============================
# SISTEMA DE CORREÇÃO DE ENTRADAS
# ===============================

class InputCorrecao:
    """Sistema que permite voltar à pergunta anterior"""
    def __init__(self):
        self.historico = []
        
    def perguntar(self, prompt, validacao=None, tipo=str, padrao=None):
        """Faz uma pergunta com possibilidade de correção"""
        while True:
            try:
                valor = input(prompt).strip()
                
                # Comando especial para voltar
                if valor.lower() == 'voltar' and self.historico:
                    return 'VOLTAR'
                
                # Aplicar conversão de tipo
                if tipo == int:
                    valor = int(valor)
                elif tipo == float:
                    valor = float(valor)
                
                # Aplicar validação personalizada
                if validacao and not validacao(valor):
                    print("Valor inválido! Tente novamente.")
                    continue
                    
                # Se chegou aqui, valor é válido
                self.historico.append((prompt, valor))
                return valor
                
            except ValueError:
                print(f"Tipo inválido! Digite um valor do tipo {tipo.__name__}.")
            except Exception as e:
                print(f"Erro: {str(e)} - Tente novamente")
    
    def voltar(self):
        """Volta para a pergunta anterior"""
        if self.historico:
            self.historico.pop()  # Remove a última entrada
            return True
        return False

# ===============================
# FUNÇÕES DE CONFIGURAÇÃO
# ===============================

def configurar_sistema():
    global CONFIG
    
    print("\n" + "="*60)
    print("🛠️ CONFIGURAÇÃO DO SISTEMA DE AUTOMAÇÃO DE DOCUMENTOS")
    print("="*60)
    print("Durante a configuração, digite 'voltar' a qualquer momento")
    print("para corrigir a informação anterior!\n")
    
    # Carregar configuração existente
    loaded_config = carregar_configuracao()
    if loaded_config:
        CONFIG.update(loaded_config)
    
    # Sistema de correção
    corretor = InputCorrecao()
    
    # Configurar diretórios
    print("\n📁 CONFIGURAÇÃO DE DIRETÓRIOS")
    while True:
        modelos = corretor.perguntar(
            "Caminho COMPLETO para a pasta de modelos: ",
            padrao=CONFIG.get('diretorios', {}).get('modelos', '')
        )
        if modelos == 'VOLTAR': 
            if not corretor.voltar(): print("Não há o que voltar!")
            continue
        CONFIG['diretorios']['modelos'] = modelos
        break
    
    while True:
        base_dados = corretor.perguntar(
            "Caminho COMPLETO para o arquivo Excel com dados: ",
            padrao=CONFIG.get('diretorios', {}).get('base_dados', '')
        )
        if base_dados == 'VOLTAR': 
            corretor.voltar()
            continue
        CONFIG['diretorios']['base_dados'] = base_dados
        break
    
    while True:
        saida = corretor.perguntar(
            "Caminho COMPLETO para salvar documentos gerados: ",
            padrao=CONFIG.get('diretorios', {}).get('saida', '')
        )
        if saida == 'VOLTAR': 
            corretor.voltar()
            continue
        CONFIG['diretorios']['saida'] = saida
        
        # Criar diretório de saída se necessário
        saida_limpo = limpar_caminho(saida)
        if not os.path.exists(saida_limpo):
            try:
                os.makedirs(saida_limpo)
                print(f"✓ Criado diretório: {saida_limpo}")
            except Exception as e:
                print(f"⚠ Erro ao criar diretório: {str(e)}")
                continue
        break
    
    # Configurar placeholders
    print("\n🔣 CONFIGURAÇÃO DE PLACEHOLDERS")
    print("Defina cada placeholder e seu significado:")
    print("IMPORTANTE: Use o NOME EXATO da coluna do Excel, como aparece na primeira linha da planilha!")
    
    # Limpar mapeamento anterior
    CONFIG['placeholders'] = {}
    
    while True:
        print("\n" + "-"*40)
        placeholder = corretor.perguntar(
            "Digite o placeholder (ex: NOME_COMPLETO) ou 'sair' para terminar: "
        )
        if placeholder == 'VOLTAR':
            if not corretor.voltar(): 
                print("Não há o que voltar!")
            continue
            
        if placeholder.lower() == 'sair':
            break
            
        if not placeholder:
            print("⚠ Placeholder não pode ser vazio!")
            continue
            
        significado = corretor.perguntar(
            f"O que '{placeholder}' representa? (ex: Coluna no Excel com nomes): "
        )
        if significado == 'VOLTAR':
            corretor.voltar()
            continue
            
        if not significado:
            print("⚠ Descrição obrigatória!")
            continue
            
        coluna = corretor.perguntar(
            f"Qual é o NOME DA COLUNA no Excel que corresponde a '{placeholder}'? "
            f"\n(Digite o nome exato como aparece na planilha, ex: 'Nome Completo', 'CPF'): "
        )
        if coluna == 'VOLTAR':
            corretor.voltar()
            continue
            
        if not coluna:
            print("⚠ Coluna obrigatória!")
            continue
            
        # Confirmar antes de salvar
        print(f"\nResumo do placeholder:")
        print(f"  Marcador: [{placeholder}]")
        print(f"  Descrição: {significado}")
        print(f"  Coluna no Excel: '{coluna}'")
        
        confirmar = perguntar_sim_nao("Confirmar este placeholder?")
        if not confirmar:
            print("Placeholder descartado. Vamos recomeçar.")
            continue
            
        CONFIG['placeholders'][placeholder] = {
            "descricao": significado,
            "coluna": coluna
        }
        print(f"✓ Mapeado: {placeholder} → {significado} ({coluna})")
    
    # Configurações avançadas
    print("\n⚙️ CONFIGURAÇÕES AVANÇADAS")
    
    while True:
        segundo_plano = corretor.perguntar(
            "Executar Office em segundo plano? [s/n] ",
            tipo=str,
            padrao='s' if CONFIG.get('config_geral', {}).get('executar_em_segundo_plano', True) else 'n'
        )
        if segundo_plano == 'VOLTAR': 
            corretor.voltar()
            continue
            
        CONFIG['config_geral']['executar_em_segundo_plano'] = segundo_plano.lower() == 's'
        break
    
    # Padrão de nomeação de arquivos
    print("\n📝 CONFIGURAÇÃO DE NOME DOS ARQUIVOS")
    print("Defina como os documentos gerados serão nomeados:")
    print("Você pode usar:")
    print("  - Placeholders: [NOME_COMPLETO], [MATRICULA], etc")
    print("  - Marcadores especiais: [CONTADOR] (número sequencial), [DATA] (data atual)")
    print("  - Nomes de colunas: [Nome da Coluna] (qualquer coluna da planilha Excel)")
    print("  - Texto livre")
    print("\nExemplo: 'Contrato_[NOME_COMPLETO]_[DATA].docx'")
    print("Exemplo: 'Documento_[Departamento]_[Matrícula].docx'")
    
    padrao_nome = corretor.perguntar(
        "Padrão de nomeação para os documentos: ",
        padrao=CONFIG.get('config_geral', {}).get('padrao_nome_arquivo', 'Documento_[CONTADOR].docx')
    )
    if padrao_nome == 'VOLTAR': 
        corretor.voltar()
    else:
        # Garantir que termina com .docx
        if not padrao_nome.lower().endswith('.docx'):
            padrao_nome += '.docx'
        CONFIG['config_geral']['padrao_nome_arquivo'] = padrao_nome
    
    # Organização de saída
    while True:
        organizar = corretor.perguntar(
            "Deseja organizar documentos em subpastas? [s/n] ",
            tipo=str,
            padrao='s' if CONFIG.get('organizacao', {}).get('ativo', False) else 'n'
        )
        if organizar == 'VOLTAR': 
            corretor.voltar()
            continue
            
        CONFIG['organizacao']['ativo'] = organizar.lower() == 's'
        
        if CONFIG['organizacao']['ativo']:
            coluna_org = corretor.perguntar(
                "Qual coluna usar para organização? (ex: Departamento): "
            )
            if coluna_org == 'VOLTAR': 
                corretor.voltar()
                continue
                
            CONFIG['organizacao']['coluna'] = coluna_org
            
            limpar_chars = corretor.perguntar(
                "Limpar caracteres inválidos em nomes de pastas? [s/n] ",
                tipo=str,
                padrao='s' if CONFIG.get('organizacao', {}).get('limpar_caracteres', True) else 'n'
            )
            if limpar_chars == 'VOLTAR': 
                corretor.voltar()
                continue
                
            CONFIG['organizacao']['limpar_caracteres'] = limpar_chars.lower() == 's'
        break
    
    # Configuração de modelo por funcionário
    print("\n📄 CONFIGURAÇÃO DE MODELO POR FUNCIONÁRIO")
    while True:
        usar_modelo_especifico = corretor.perguntar(
            "Deseja usar modelos diferentes para cada funcionário? [s/n] ",
            tipo=str,
            padrao='n'
        )
        if usar_modelo_especifico == 'VOLTAR': 
            corretor.voltar()
            continue
            
        if usar_modelo_especifico.lower() == 's':
            coluna_modelo = corretor.perguntar(
                "Qual coluna contém o nome do modelo? (ex: Modelo): "
            )
            if coluna_modelo == 'VOLTAR': 
                corretor.voltar()
                continue
            CONFIG['modelo_especifico'] = {
                'ativo': True,
                'coluna': coluna_modelo
            }
        else:
            CONFIG['modelo_especifico'] = {'ativo': False}
        break
    
    # Configuração de log
    print("\n📝 CONFIGURAÇÃO DE LOGS")
    placeholder_log = corretor.perguntar(
        "Qual placeholder representa o nome do funcionário (para logs)? (ex: NOME_COMPLETO): ",
        padrao=CONFIG.get('placeholder_log', '')
    )
    if placeholder_log == 'VOLTAR': 
        corretor.voltar()
    else:
        CONFIG['placeholder_log'] = placeholder_log
    
    # Salvar configuração
    salvar_configuracao()
    print("\n✅ Configuração concluída! Pressione Enter para continuar.")
    input()

# ===============================
# FUNÇÕES DE PROCESSAMENTO COM python-docx
# ===============================

def substituir_texto_com_docx(modelo_path, substituicoes, caminho_completo):
    """Substitui placeholders usando python-docx preservando formatação"""
    try:
        # Carregar documento
        doc = Document(modelo_path)
        
        # Ordenar placeholders do maior para o menor para evitar substituições parciais
        sorted_ph = sorted(substituicoes.keys(), key=len, reverse=True)
        
        # Função para substituir em um parágrafo
        def substituir_no_paragrafo(paragraph):
            for ph in sorted_ph:
                if ph in paragraph.text:
                    valor = substituicoes[ph]
                    # Preservar formatação original
                    for run in paragraph.runs:
                        if ph in run.text:
                            run.text = run.text.replace(ph, valor)
        
        # Função para substituir em tabelas
        def substituir_em_tabelas():
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            substituir_no_paragrafo(paragraph)
        
        # Função para substituir em cabeçalhos e rodapés
        def substituir_em_secoes(section):
            for paragraph in section.header.paragraphs:
                substituir_no_paragrafo(paragraph)
            for paragraph in section.footer.paragraphs:
                substituir_no_paragrafo(paragraph)
        
        # Processar todos os parágrafos principais
        for paragraph in doc.paragraphs:
            substituir_no_paragrafo(paragraph)
        
        # Processar tabelas
        substituir_em_tabelas()
        
        # Processar cabeçalhos e rodapés
        for section in doc.sections:
            substituir_em_secoes(section)
        
        # Salvar documento
        doc.save(caminho_completo)
        return True
        
    except Exception as e:
        print(f"⚠ Erro durante substituição com python-docx: {str(e)}")
        traceback.print_exc()
        return False

def mostrar_barra_progresso(atual, total, largura=50):
    percentual = atual / total
    completos = int(largura * percentual)
    restantes = largura - completos
    
    barra = '[' + '=' * completos + ' ' * restantes + ']'
    percentual_texto = f"{percentual:.1%}"
    
    sys.stdout.write(f"\r{barra} {atual}/{total} ({percentual_texto})")
    sys.stdout.flush()

def validar_dados(df):
    """Realiza validação avançada dos dados antes do processamento"""
    problemas = []
    
    # Verificar colunas necessárias
    colunas_necessarias = [info['coluna'] for info in CONFIG['placeholders'].values()]
    colunas_faltantes = [col for col in colunas_necessarias if col not in df.columns]
    
    if colunas_faltantes:
        problemas.append(f"Colunas faltando: {', '.join(colunas_faltantes)}")
    
    # Verificar valores em branco em colunas críticas
    for ph, info in CONFIG['placeholders'].items():
        coluna = info['coluna']
        if coluna in df.columns:
            nulos = df[coluna].isnull().sum()
            if nulos > 0:
                problemas.append(f"{nulos} valores nulos na coluna '{coluna}' (para {ph})")
    
    # Verificar duplicatas em colunas chave
    if CONFIG['organizacao'].get('ativo', False):
        col_org = CONFIG['organizacao']['coluna']
        if col_org in df.columns and df[col_org].duplicated().any():
            problemas.append(f"Valores duplicados na coluna de organização '{col_org}'")
    
    return problemas

def gerar_nome_arquivo(registro, idx, cabecalhos):
    """Gera o nome do arquivo baseado no padrão configurado usando dados da planilha"""
    padrao = CONFIG['config_geral'].get('padrao_nome_arquivo', 'Documento_[CONTADOR].docx')
    
    # Substituir marcadores especiais
    nome_arquivo = padrao
    nome_arquivo = nome_arquivo.replace('[CONTADOR]', str(idx))
    nome_arquivo = nome_arquivo.replace('[DATA]', datetime.now().strftime('%d/%m/%Y'))
    nome_arquivo = nome_arquivo.replace('[HORA]', datetime.now().strftime('%H%M%S'))
    
    # Substituir por valores das colunas da planilha
    for coluna in cabecalhos:
        if coluna in registro:
            # Usar colchetes como marcadores: [Nome da Coluna]
            marcador = f'[{coluna}]'
            if marcador in nome_arquivo:
                valor = registro[coluna]
                
                # Formatar datas corretamente
                if isinstance(valor, pd.Timestamp) or isinstance(valor, datetime):
                    valor = valor.strftime('%d/%m/%Y')
                else:
                    valor = str(valor)
                
                nome_arquivo = nome_arquivo.replace(marcador, valor)
    
    # Remover caracteres inválidos
    nome, ext = os.path.splitext(nome_arquivo)
    nome_limpo = limpar_nome_arquivo(nome)
    
    return nome_limpo + ext

def processar_documento_individual(modelo_path, caminho_completo, subs):
    """Função usando python-docx para substituição"""
    try:
        # Verificar se existem placeholders para substituir
        if not subs:
            print("⚠ Nenhum placeholder para substituir! Verifique o mapeamento.")
            return False

        # Usar python-docx para substituição
        sucesso = substituir_texto_com_docx(modelo_path, subs, caminho_completo)
        return sucesso
        
    except Exception as e:
        print(f"❌ Erro crítico ao processar documento: {str(e)}")
        traceback.print_exc()
        return False

def normalizar_nome(nome):
    """Normaliza nomes removendo acentos, espaços e caracteres especiais"""
    # Remover acentos
    nome = unicodedata.normalize('NFKD', nome).encode('ascii', 'ignore').decode('ascii')
    # Converter para minúsculas
    nome = nome.lower()
    # Remover espaços e caracteres especiais
    nome = re.sub(r'[\s_\-]+', '', nome)
    return nome

def encontrar_modelo(nome_modelo, modelos_por_nome, modelos_por_base_sem_ext, modelos_por_nome_normalizado):
    """Busca inteligente por modelos com diferentes estratégias"""
    # 1. Tentar nome exato (com extensão)
    if nome_modelo in modelos_por_nome:
        return modelos_por_nome[nome_modelo]
    
    # 2. Tentar nome exato em minúsculas
    nome_lower = nome_modelo.lower()
    for nome, path in modelos_por_nome.items():
        if nome.lower() == nome_lower:
            return path
    
    # 3. Tentar adicionar extensões comuns
    for ext in ['.docx', '.doc']:
        nome_tentativa = nome_modelo + ext
        if nome_tentativa in modelos_por_nome:
            return modelos_por_nome[nome_tentativa]
    
    # 4. Tentar sem extensão (nome base)
    base_sem_ext = os.path.splitext(nome_modelo)[0]
    if base_sem_ext in modelos_por_base_sem_ext:
        return modelos_por_base_sem_ext[base_sem_ext]
    
    # 5. Tentar sem extensão em minúsculas
    base_sem_ext_lower = base_sem_ext.lower()
    for base, path in modelos_por_base_sem_ext.items():
        if base.lower() == base_sem_ext_lower:
            return path
    
    # 6. Tentar combinações de caminhos
    for path in modelos_por_nome.values():
        if nome_modelo.lower() in path.lower():
            return path
    
    # 7. Tentar versão normalizada (sem acentos, espaços e caracteres especiais)
    nome_normalizado = normalizar_nome(nome_modelo)
    if nome_normalizado in modelos_por_nome_normalizado:
        return modelos_por_nome_normalizado[nome_normalizado]
    
    # 8. Tentar versão normalizada sem extensão
    base_normalizado = normalizar_nome(base_sem_ext)
    if base_normalizado in modelos_por_nome_normalizado:
        return modelos_por_nome_normalizado[base_normalizado]
    
    return None

def processar_documentos():
    try:
        print("\n" + "="*60)
        print("🚀 INICIANDO PROCESSAMENTO DE DOCUMENTOS")
        print("="*60)
        
        # Verificar configuração mínima
        if not CONFIG['placeholders']:
            print("❌ Nenhum placeholder configurado! Execute a configuração primeiro.")
            input("\nPressione Enter para voltar...")
            return
            
        # Carregar base de dados
        try:
            # Obter e limpar caminho
            caminho_base = limpar_caminho(CONFIG['diretorios']['base_dados'])
            
            # Verificar se arquivo existe
            if not os.path.exists(caminho_base):
                print(f"❌ Arquivo não encontrado: {caminho_base}")
                input("\nPressione Enter para voltar...")
                return
                
            # Ler apenas os cabeçalhos para validação
            df_cabecalhos = pd.read_excel(caminho_base, nrows=0)
            cabecalhos = list(df_cabecalhos.columns)
            print(f"✓ Cabeçalhos encontrados na planilha: {', '.join(cabecalhos)}")
            
            # Verificar colunas necessárias
            colunas_faltantes = []
            for ph, info in CONFIG['placeholders'].items():
                coluna = info['coluna']
                if coluna not in cabecalhos:
                    colunas_faltantes.append(f"'{coluna}' (para placeholder {ph})")
            
            if colunas_faltantes:
                print("❌ Colunas faltando na base de dados:")
                for col in colunas_faltantes:
                    print(f"  - {col}")
                input("\nPressione Enter para voltar...")
                return
                
            # Carregar dados completos se todas colunas existirem
            df = pd.read_excel(caminho_base)
            total_registros = len(df)
            print(f"✓ Base de dados carregada: {total_registros} registros encontrados")
            
        except Exception as e:
            print(f"❌ Erro ao carregar base de dados: {str(e)}")
            traceback.print_exc()
            input("\nPressione Enter para voltar...")
            return
            
        # Verificar modelos
        caminho_modelos = limpar_caminho(CONFIG['diretorios']['modelos'])
        print(f"🔍 Verificando acesso ao diretório de modelos: {caminho_modelos}")
        
        # Verificar permissões
        if not os.access(caminho_modelos, os.R_OK):
            print(f"❌ Sem permissão de leitura no diretório: {caminho_modelos}")
            input("\nPressione Enter para voltar...")
            return
            
        modelos = []
        print(f"\n🔍 Procurando modelos em: {caminho_modelos}")
        for root, _, files in os.walk(caminho_modelos):
            for file in files:
                if file.lower().endswith(('.docx', '.doc')):
                    full_path = os.path.join(root, file)
                    modelos.append(full_path)
                    print(f"  ✅ Encontrado: {file}")
        
        if not modelos:
            print("❌ Nenhum modelo Word encontrado!")
            input("\nPressione Enter para voltar...")
            return
            
        print(f"✓ {len(modelos)} modelos encontrados")
        
        # Criar estruturas para busca eficiente de modelos
        modelos_por_nome = {}
        modelos_por_base_sem_ext = {}
        modelos_por_nome_normalizado = {}
        
        for modelo_path in modelos:
            nome_arquivo = os.path.basename(modelo_path)
            modelos_por_nome[nome_arquivo] = modelo_path
            
            # Criar entrada sem extensão
            base_sem_ext = os.path.splitext(nome_arquivo)[0]
            modelos_por_base_sem_ext[base_sem_ext] = modelo_path
            
            # Criar versão normalizada (sem acentos, espaços e caracteres especiais)
            nome_normalizado = normalizar_nome(nome_arquivo)
            base_normalizado = normalizar_nome(base_sem_ext)
            modelos_por_nome_normalizado[nome_normalizado] = modelo_path
            modelos_por_nome_normalizado[base_normalizado] = modelo_path
        
        # Se for caminho de rede, tentar mapear unidade (apenas para Windows)
        if caminho_modelos.startswith('\\\\'):
            print("ⓘ Caminho de rede detectado, ativando compatibilidade...")
            # Tenta mapear a unidade de rede (apenas Windows)
            try:
                drive_letter = 'Z:'  # Pode ser necessário escolher uma letra livre
                os.system(f'net use {drive_letter} {caminho_modelos} /persistent:yes > nul 2>&1')
                caminho_modelos = drive_letter + '\\'
                print(f"✓ Caminho de rede mapeado para {drive_letter}")
            except:
                print("⚠ Não foi possível mapear o caminho de rede. Continuando...")
        
        # Sistema de checkpoint
        saida_path = limpar_caminho(CONFIG['diretorios']['saida'])
        checkpoint_file = os.path.join(saida_path, 'checkpoint.json')
        checkpoint = {}
        
        # Carregar checkpoint se existir
        if os.path.exists(checkpoint_file):
            try:
                with open(checkpoint_file, 'r') as f:
                    checkpoint = json.load(f)
                print(f"✓ Checkpoint encontrado: Continuando do registro {checkpoint.get('ultimo_registro', 0)}")
            except:
                print("⚠ Erro ao carregar checkpoint, iniciando do zero")
        
        # Processar cada registro
        total_processados = 0
        erros = []  # Lista de dicionários com detalhes de erros
        categorias = set()
        modelos_faltantes = {}  # Dicionário para rastrear modelos faltantes
        
        print("\n⏳ Gerando documentos...")
        inicio = time.time()
        
        # Mostrar barra de progresso inicial
        mostrar_barra_progresso(0, total_registros)
        
        # Índice inicial
        start_idx = checkpoint.get('ultimo_registro', 0)
        
        for idx, (index, registro) in enumerate(df.iterrows(), 1):
            # Pular registros já processados
            if idx <= start_idx:
                mostrar_barra_progresso(idx, total_registros)
                continue
            
            try:
                # Obter nome do funcionário para logs
                nome_funcionario = "Desconhecido"
                if 'placeholder_log' in CONFIG and CONFIG['placeholder_log'] in CONFIG['placeholders']:
                    coluna_log = CONFIG['placeholders'][CONFIG['placeholder_log']]['coluna']
                    if coluna_log in registro:
                        nome_funcionario = str(registro[coluna_log])
                
                # CORREÇÃO: Tratamento robusto para valores nulos/vazios
                subs = {}
                for ph, info in CONFIG['placeholders'].items():
                    coluna = info['coluna']
                    valor = registro[coluna]
                    
                    # Tratamento para valores ausentes/inválidos
                    if pd.isna(valor) or valor is None:
                        valor_formatado = ""  # Valor vazio para campos ausentes
                    else:
                        # Verificar se é uma data e formatar corretamente
                        if isinstance(valor, (pd.Timestamp, datetime)):
                            try:
                                valor_formatado = valor.strftime('%d/%m/%Y')
                            except (ValueError, AttributeError):
                                valor_formatado = ""  # Fallback para datas inválidas
                        else:
                            valor_formatado = str(valor)
                    
                    subs[ph] = valor_formatado
                
                # Selecionar modelo apropriado
                modelo_path = None
                if CONFIG.get('modelo_especifico', {}).get('ativo', False):
                    coluna_modelo = CONFIG['modelo_especifico']['coluna']
                    nome_modelo = str(registro[coluna_modelo])
                    
                    # Usar sistema inteligente de busca
                    modelo_path = encontrar_modelo(
                        nome_modelo, 
                        modelos_por_nome, 
                        modelos_por_base_sem_ext,
                        modelos_por_nome_normalizado
                    )
                    
                    if not modelo_path:
                        # Registrar modelo faltante com contagem
                        if nome_modelo in modelos_faltantes:
                            modelos_faltantes[nome_modelo]['contagem'] += 1
                        else:
                            modelos_faltantes[nome_modelo] = {
                                'contagem': 1,
                                'funcionarios': set()
                            }
                        modelos_faltantes[nome_modelo]['funcionarios'].add(nome_funcionario)
                        
                        erros.append({
                            'indice': idx,
                            'nome': nome_funcionario,
                            'erro': f"Modelo '{nome_modelo}' não encontrado",
                            'modelo': nome_modelo
                        })
                        print(f"\n❌ Modelo não encontrado: '{nome_modelo}' para {nome_funcionario}")
                        continue
                else:
                    modelo_path = modelos[0]  # Usar primeiro modelo
                
                # Organizar por categoria se necessário
                saida_path_atual = saida_path
                if CONFIG['organizacao'].get('ativo', False):
                    categoria = str(registro[CONFIG['organizacao']['coluna']])
                    
                    if CONFIG['organizacao'].get('limpar_caracteres', False):
                        categoria = limpar_nome_arquivo(categoria)
                    
                    categoria_path = os.path.join(saida_path_atual, categoria)
                    
                    if categoria not in categorias:
                        try:
                            os.makedirs(categoria_path, exist_ok=True)
                            categorias.add(categoria)
                        except Exception as e:
                            print(f"⚠ Erro ao criar pasta {categoria}: {str(e)}")
                            # Usar diretório principal em caso de erro
                            continue
                    
                    saida_path_atual = categoria_path
                
                # Gerar nome de arquivo personalizado usando dados da planilha
                nome_arquivo = gerar_nome_arquivo(registro, idx, cabecalhos)
                caminho_completo = os.path.join(saida_path_atual, nome_arquivo)
                
                # Processar documento individual
                sucesso = processar_documento_individual(modelo_path, caminho_completo, subs)
                
                if sucesso:
                    total_processados += 1
                    print(f"\n✓ Documento gerado para {nome_funcionario}: {nome_arquivo}")
                
                    # Salvar checkpoint após cada documento processado com sucesso
                    checkpoint = {'ultimo_registro': idx}
                    try:
                        with open(checkpoint_file, 'w') as f:
                            json.dump(checkpoint, f)
                    except Exception as e:
                        print(f"⚠ Erro ao salvar checkpoint: {str(e)}")
                else:
                    erros.append({
                        'indice': idx,
                        'nome': nome_funcionario,
                        'erro': "Falha ao gerar documento",
                        'modelo': modelo_path
                    })
                    print(f"\n❌ Falha ao gerar documento para {nome_funcionario}")
                
            except Exception as e:
                # Registrar erro com detalhes
                nome_funcionario = "Desconhecido"
                if 'placeholder_log' in CONFIG and CONFIG['placeholder_log'] in CONFIG['placeholders']:
                    coluna_log = CONFIG['placeholders'][CONFIG['placeholder_log']]['coluna']
                    if coluna_log in registro:
                        nome_funcionario = str(registro[coluna_log])
                
                erros.append({
                    'indice': idx,
                    'nome': nome_funcionario,
                    'erro': str(e),
                    'modelo': modelo_path if 'modelo_path' in locals() else "Não definido"
                })
                print(f"\n❌ Erro no registro {idx} ({nome_funcionario}): {str(e)}")
                traceback.print_exc()
                
                # Pausa para evitar sobrecarga
                time.sleep(2)
            
            # Atualizar barra de progresso
            mostrar_barra_progresso(idx, total_registros)
        
        # Remover checkpoint após conclusão
        if os.path.exists(checkpoint_file):
            try:
                os.remove(checkpoint_file)
            except:
                pass
        
        tempo_total = time.time() - inicio
        print(f"\n\n✅ Processamento concluído em {tempo_total:.1f} segundos")
        
        # Relatório final
        print("\n" + "="*50)
        print(f"📊 RESUMO DO PROCESSAMENTO")
        print("="*50)
        print(f"• Documentos gerados: {total_processados}/{total_registros}")
        print(f"• Erros encontrados: {len(erros)}")
        if tempo_total > 0:
            print(f"• Velocidade: {total_processados/tempo_total:.1f} docs/segundo")
        
        # Salvar relatório detalhado
        log_path = os.path.join(saida_path, "relatorio_geracao.txt")
        try:
            with open(log_path, 'w', encoding='utf-8') as f:
                f.write("RELATÓRIO DE GERAÇÃO DE DOCUMENTOS\n")
                f.write("="*50 + "\n\n")
                f.write(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")
                f.write(f"Total de registros: {total_registros}\n")
                f.write(f"Documentos gerados: {total_processados}\n")
                f.write(f"Documentos com erro: {len(erros)}\n")
                f.write(f"Tempo total: {tempo_total:.1f} segundos\n")
                if tempo_total > 0:
                    f.write(f"Velocidade média: {total_processados/tempo_total:.1f} docs/segundo\n\n")
                
                f.write("CONFIGURAÇÃO UTILIZADA:\n")
                f.write(f"• Pasta de modelos: {CONFIG['diretorios']['modelos']}\n")
                f.write(f"• Arquivo de dados: {CONFIG['diretorios']['base_dados']}\n")
                f.write(f"• Pasta de saída: {CONFIG['diretorios']['saida']}\n")
                f.write(f"• Padrão de nome: {CONFIG['config_geral']['padrao_nome_arquivo']}\n\n")
                
                f.write("PLACEHOLDERS UTILIZADOS:\n")
                for ph, info in CONFIG['placeholders'].items():
                    f.write(f"- {ph}: {info['descricao']} ({info['coluna']})\n")
                
                if erros:
                    f.write("\n\nERROS DETALHADOS:\n")
                    f.write("="*50 + "\n")
                    for erro in erros:
                        f.write(f"\n• REGISTRO #{erro['indice']}\n")
                        f.write(f"  Nome: {erro['nome']}\n")
                        f.write(f"  Modelo: {erro.get('modelo', 'Não especificado')}\n")
                        f.write(f"  Erro: {erro['erro']}\n")
                        f.write("-"*50 + "\n")
                else:
                    f.write("\n\nNENHUM ERRO ENCONTRADO DURANTE O PROCESSAMENTO\n")
                
                # Seção detalhada de modelos faltantes
                if modelos_faltantes:
                    f.write("\n\n🔍 MODELOS FALTANTES DETECTADOS:\n")
                    f.write("="*50 + "\n")
                    f.write(f"Total de modelos faltantes: {len(modelos_faltantes)}\n")
                    f.write(f"Total de ocorrências: {sum(info['contagem'] for info in modelos_faltantes.values())}\n\n")
                    
                    for modelo, info in modelos_faltantes.items():
                        f.write(f"\n• MODELO: {modelo}\n")
                        f.write(f"  Ocorrências: {info['contagem']}\n")
                        f.write(f"  Funcionários afetados ({min(5, len(info['funcionarios']))} de {len(info['funcionarios'])}):\n")
                        
                        # Listar até 5 funcionários afetados
                        for i, func in enumerate(list(info['funcionarios'])[:5]):
                            f.write(f"    - {func}\n")
                        
                        if len(info['funcionarios']) > 5:
                            f.write(f"    ... e mais {len(info['funcionarios']) - 5}\n")
                        
                        f.write("-"*50 + "\n")
                    
                    f.write("\nMODELOS DISPONÍVEIS NA PASTA:\n")
                    f.write("="*50 + "\n")
                    for modelo in modelos_por_nome.keys():
                        f.write(f"- {modelo}\n")
            
            print(f"\n📝 Relatório completo salvo em: {log_path}")
        except Exception as e:
            print(f"⚠ Não foi possível salvar relatório: {str(e)}")
            
    except Exception as e:
        print(f"\n❌ ERRO GRAVE: {str(e)}")
        traceback.print_exc()
    finally:
        input("\nPressione Enter para voltar ao menu...")

# ===============================
# MENU PRINCIPAL
# ===============================

def menu_principal():
    try:
        loaded_config = carregar_configuracao()
        if loaded_config:
            CONFIG.update(loaded_config)
    except Exception as e:
        print(f"⚠ Erro ao carregar configuração: {str(e)}")
        print("Iniciando com configuração padrão.")
    
    while True:
        print("\n" + "="*60)
        print("🤖 DOCUMENT AUTOMATOR - SISTEMA DE AUTOMAÇÃO")
        print("="*60)
        print("1. Configurar sistema")
        print("2. Processar documentos")
        print("3. Visualizar configuração atual")
        print("4. Sair")
        
        opcao = input("\nSelecione uma opção: ").strip()
        
        if opcao == '1':
            configurar_sistema()
        elif opcao == '2':
            if not CONFIG.get('diretorios') or not CONFIG.get('placeholders'):
                print("⚠ Configure o sistema primeiro!")
                input("Pressione Enter para continuar...")
            else:
                processar_documentos()
        elif opcao == '3':
            print("\nConfiguração atual:")
            print(json.dumps(CONFIG, indent=4, ensure_ascii=False))
            print(f"\nArquivo de configuração: {CONFIG_FILE}")
            input("\nPressione Enter para voltar...")
        elif opcao == '4':
            print("\n✅ Sistema encerrado. Até logo!")
            break
        else:
            print("⚠ Opção inválida! Tente novamente.")
            time.sleep(1)

# ===============================
# PONTO DE ENTRADA
# ===============================

if __name__ == "__main__":
    try:
        menu_principal()
    except Exception as e:
        print(f"\n❌ ERRO INESPERADO: {str(e)}")
        traceback.print_exc()
        input("\nPressione Enter para sair...")
    finally:
        # Garantir que o Word seja fechado
        try:
            word = win32.Dispatch("Word.Application")
            word.Quit()
        except:
            pass